#!/usr/bin/env python3
"""Render a Supra Markdown document to Word (.docx).

Produces a real Word document -- Heading styles, native tables, a page-number
footer -- so the output is editable and can carry review comments, which is what
audit and review documents actually get used for.

Supports the Markdown subset used by the Supra docs: ATX headings, pipe tables,
fenced code blocks, blockquotes, bullet/numbered lists, horizontal rules, and
inline **bold**, `code` and [links](url).

Usage:
    python3 scripts/md2docx.py <input.md> <output.docx> [--title T] [--subtitle S]
"""

import argparse
import re
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor

# Brand palette, matching the HTML renderings in docs/.
NAVY = RGBColor(0x1A, 0x23, 0x7E)
NAVY_HEX = "1A237E"
H2_COLOR = RGBColor(0x28, 0x35, 0x93)
H3_COLOR = RGBColor(0x39, 0x49, 0xAB)
CODE_RED = RGBColor(0xC6, 0x28, 0x28)
GREY_HEX = "F2F2F2"
AMBER_HEX = "FFF8E1"

MONO = "Consolas"
BODY = "Calibri"


# --------------------------------------------------------------------------- #
# low-level helpers
# --------------------------------------------------------------------------- #
def shade(element, hex_fill):
    """Apply cell/paragraph background shading."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    element.append(shd)


def set_cell_background(cell, hex_fill):
    shade(cell._tc.get_or_add_tcPr(), hex_fill)


def repeat_header_row(row):
    """Mark a table row as a header so it repeats across page breaks."""
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_page_number_footer(section):
    para = section.footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("Page ")
    run.font.size = Pt(8)
    run.font.name = BODY

    for instr in ("PAGE", "NUMPAGES"):
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), instr)
        para._p.append(fld)
        if instr == "PAGE":
            sep = para.add_run(" of ")
            sep.font.size = Pt(8)
            sep.font.name = BODY


# --------------------------------------------------------------------------- #
# inline formatting
# --------------------------------------------------------------------------- #
INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")


def add_inline(paragraph, text):
    """Add text to a paragraph, honouring **bold**, `code` and [links](url)."""
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = MONO
            run.font.size = Pt(9)
            run.font.color.rgb = CODE_RED
        elif part.startswith("["):
            label, _ = part[1:].split("](", 1)
            # Rendered as plain emphasised text: the TOC targets are internal
            # anchors, which Word would not resolve anyway.
            run = paragraph.add_run(label)
            run.font.color.rgb = NAVY
        else:
            paragraph.add_run(part)


# --------------------------------------------------------------------------- #
# block builders
# --------------------------------------------------------------------------- #
def add_heading(doc, text, level):
    para = doc.add_heading(level=level)
    color = {1: NAVY, 2: H2_COLOR, 3: H3_COLOR}[level]
    size = {1: 20, 2: 15, 3: 12.5}[level]
    add_inline(para, text)
    for run in para.runs:
        run.font.color.rgb = color
        run.font.name = BODY
        run.font.size = Pt(size)
        run.bold = True
    return para


def add_code_block(doc, lines):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.15)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    shade(para._p.get_or_add_pPr(), GREY_HEX)
    run = para.add_run("\n".join(lines))
    run.font.name = MONO
    run.font.size = Pt(8)
    return para


def add_blockquote(doc, lines):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.3)
    para.paragraph_format.right_indent = Inches(0.3)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    shade(para._p.get_or_add_pPr(), AMBER_HEX)
    add_inline(para, " ".join(lines))
    for run in para.runs:
        run.font.size = Pt(9.5)
    return para


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def add_table(doc, rows):
    header, body = rows[0], rows[1:]
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for cell, text in zip(table.rows[0].cells, header):
        set_cell_background(cell, NAVY_HEX)
        para = cell.paragraphs[0]
        add_inline(para, text)
        if not para.runs:  # empty header cell
            para.add_run("")
        for run in para.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9)
    repeat_header_row(table.rows[0])

    for values in body:
        cells = table.add_row().cells
        # tolerate ragged rows rather than dropping content
        for i, text in enumerate(values[: len(header)]):
            para = cells[i].paragraphs[0]
            add_inline(para, text)
            for run in para.runs:
                run.font.size = Pt(9)
    return table


# --------------------------------------------------------------------------- #
# document conversion
# --------------------------------------------------------------------------- #
def convert(md_path, docx_path, title, subtitle):
    lines = open(md_path, encoding="utf-8").read().splitlines()

    doc = Document()

    section = doc.sections[0]
    section.page_width, section.page_height = Mm(210), Mm(297)  # A4
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Mm(20))
    add_page_number_footer(section)

    normal = doc.styles["Normal"]
    normal.font.name = BODY
    normal.font.size = Pt(10.5)

    if title:
        head = doc.add_paragraph()
        head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = head.add_run(title)
        run.bold = True
        run.font.size = Pt(24)
        run.font.color.rgb = NAVY
        if subtitle:
            sub = doc.add_paragraph()
            sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            srun = sub.add_run(subtitle)
            srun.font.size = Pt(10)
            srun.font.color.rgb = H3_COLOR

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # fenced code block
        if stripped.startswith("```"):
            i += 1
            block = []
            while i < n and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1
            add_code_block(doc, block)
            continue

        # pipe table: a header line followed by a |---| separator
        if (
            stripped.startswith("|")
            and i + 1 < n
            and re.match(r"^\|[\s:|-]+\|$", lines[i + 1].strip())
        ):
            rows = [split_row(stripped)]
            i += 2
            while i < n and lines[i].strip().startswith("|"):
                rows.append(split_row(lines[i].strip()))
                i += 1
            add_table(doc, rows)
            doc.add_paragraph()
            continue

        # blockquote
        if stripped.startswith(">"):
            block = []
            while i < n and lines[i].strip().startswith(">"):
                block.append(lines[i].strip().lstrip(">").strip())
                i += 1
            add_blockquote(doc, [b for b in block if b])
            continue

        # headings
        m = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if m:
            add_heading(doc, m.group(2), len(m.group(1)))
            i += 1
            continue

        # horizontal rule -> page break, which is what the section dividers mean
        if re.match(r"^-{3,}$", stripped):
            doc.add_page_break()
            i += 1
            continue

        # bullet list
        if re.match(r"^[-*]\s+", stripped):
            para = doc.add_paragraph(style="List Bullet")
            add_inline(para, re.sub(r"^[-*]\s+", "", stripped))
            i += 1
            continue

        # numbered list
        if re.match(r"^\d+\.\s+", stripped):
            para = doc.add_paragraph(style="List Number")
            add_inline(para, re.sub(r"^\d+\.\s+", "", stripped))
            i += 1
            continue

        # blank line
        if not stripped:
            i += 1
            continue

        # paragraph: join wrapped lines until a blank or a new block starts
        block = []
        while i < n:
            cur = lines[i].strip()
            if not cur or re.match(r"^(#{1,3}\s|[-*]\s|\d+\.\s|>|```|\|)", cur) \
                    or re.match(r"^-{3,}$", cur):
                break
            block.append(cur)
            i += 1
        para = doc.add_paragraph()
        text = " ".join(block)
        if text.startswith("*") and text.endswith("*") and not text.startswith("**"):
            run = para.add_run(text.strip("*"))
            run.italic = True
        else:
            add_inline(para, text)

    doc.save(docx_path)
    return docx_path


def main():
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("input")
    ap.add_argument("output")
    ap.add_argument("--title", default="")
    ap.add_argument("--subtitle", default="")
    args = ap.parse_args()

    out = convert(args.input, args.output, args.title, args.subtitle)
    print(f"DOCX generated: {out}")


if __name__ == "__main__":
    sys.exit(main())
